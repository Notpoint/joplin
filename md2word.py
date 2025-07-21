import re, urllib.parse
from pathlib import Path
from docx import Document
from docx.shared import Inches
import cairosvg  # <--- 1. 导入新安装的库

# ======== 需自行修改的路径 ========
MD_FOLDER        = Path(r"C:\Users\13684\Desktop\jophin\joplindirectoutput\航的笔记本")           # .md 所在目录
OUTPUT_FOLDER    = Path(r"C:\Users\13684\Desktop\jophin\joplin_word")      # .docx 输出目录
RESOURCE_FOLDER  = Path(r"C:\Users\13684\Desktop\jophin\joplindirectoutput\_resources")       # 图片 / 音频 目录
TRANSCRIPT_FOLDER= Path(r"C:\Users\13684\Desktop\jophin\joplin_audio2txt")        # .txt 转录文件目录
# ==================================

IMG_RGX   = re.compile(r'!\[.*?]\((.*?)\)')
AUDIO_RGX = re.compile(r'\[[^\]]+?\.m4a]\((.*?)\.m4a\)')
HTML_IMG_RGX = re.compile(r'<img src="(.*?)"')

def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    for line in md_path.read_text(encoding='utf-8').splitlines():
        # ----- 图片 -----
        md_img_m = IMG_RGX.search(line)
        html_img_m = HTML_IMG_RGX.search(line)

        if md_img_m or html_img_m:
            # 根据匹配结果获取图片链接和前缀文本
            if md_img_m:
                img_link = md_img_m.group(1).strip()
                prefix = IMG_RGX.sub('', line).strip()
            else: # html_img_m is True
                img_link = html_img_m.group(1).strip()
                prefix = HTML_IMG_RGX.sub('', line).strip()

            if prefix:
                doc.add_paragraph(prefix)

            # Joplin的资源路径可能是相对的，例如 "../_resources/xxx.png"
            if img_link.startswith('../'):
                img_link = img_link[3:] # 去掉 "../"

            # ==================== 主要修改部分 开始 ====================
            # 从链接中获取文件名部分，例如 "绘制中"
            image_filename = Path(img_link).name

            # 检查文件名本身是否包含后缀
            if not Path(image_filename).suffix:
                # 如果没有后缀，则默认它是.svg文件，并为其添加.svg后缀
                final_filename = image_filename + '.svg'
            else:
                # 如果已经有后缀，则直接使用
                final_filename = image_filename
            
            # 使用处理过的新文件名构建最终的图片文件路径
            img_file = (RESOURCE_FOLDER / final_filename).resolve()
            # ==================== 主要修改部分 结束 ====================

            if img_file.exists():
                # 判断文件是否为SVG
                if img_file.suffix.lower() == '.svg':
                    # 如果是SVG，先将其转换为PNG
                    png_path = img_file.with_suffix('.png')
                    print(f"    正在转换SVG -> PNG: {img_file.name} -> {png_path.name}")
                    cairosvg.svg2png(url=str(img_file), write_to=str(png_path))
                    # 将要插入的图片路径指向转换后的PNG文件
                    path_to_insert = png_path
                else:
                    # 如果不是SVG（是PNG, JPG等），直接使用原路径
                    path_to_insert = img_file
                
                # 使用最终确定的图片路径插入图片
                doc.add_picture(str(path_to_insert), width=Inches(4))
            else:
                # 提示信息使用我们尝试查找的文件名，更便于排查问题
                doc.add_paragraph(f'[缺失图片: {img_file.name}]')
            continue

        # ----- 音频链接 -----
        aud_m = AUDIO_RGX.search(line)
        if aud_m:
            # ① 原行照写
            doc.add_paragraph(line)

            # ② 找转录文本
            enc_name = Path(aud_m.group(1)).name + '.m4a'    # 带 %20 的文件名
            true_name = urllib.parse.unquote(enc_name)       # 解码 %20 -> 空格
            txt_file  = TRANSCRIPT_FOLDER / (Path(true_name).stem + '.txt')

            if txt_file.exists():
                text = txt_file.read_text(encoding='utf-8').strip()
            else:
                text = f'[未找到转录文件: {txt_file.name}]'

            doc.add_paragraph(text)
            continue

        # ----- 普通文本 -----
        doc.add_paragraph(line)

    doc.save(docx_path)

def batch_convert():
    OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)
    md_files = list(MD_FOLDER.glob('*.md'))
    if not md_files:
        print(f'⚠ 未在 {MD_FOLDER} 发现 .md 文件')
        return

    for md_fp in md_files:
        print(f'→ 正在处理: {md_fp.name}')
        out_fp = OUTPUT_FOLDER / (md_fp.stem + '.docx')
        try:
            md_to_docx(md_fp, out_fp)
            print(f'✓ 已生成 {out_fp.name}')
        except Exception as e:
            print(f'✗ 处理 {md_fp.name} 失败：{e}')
            # 打印更详细的错误信息，便于调试
            import traceback
            traceback.print_exc()

if __name__ == '__main__':
    batch_convert()