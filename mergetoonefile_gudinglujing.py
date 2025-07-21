import os
import glob
import datetime
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt

def auto_merge_folder_files():
    """
    自动合并指定文件夹中的所有docx文件（无GUI选择）。
    """
    
    # --- 【修改部分】硬编码文件夹路径 ---
    # 请根据您的实际路径修改以下路径
    target_folder = r"C:\Users\13684\Desktop\jophin\joplin_word"  # 修改为您的实际路径
    
    print(f"🔍 正在扫描文件夹：{target_folder}")
    
    # 检查文件夹是否存在
    if not os.path.exists(target_folder):
        print(f"❌ 错误：指定的文件夹不存在: {target_folder}")
        print("请修改代码中的 target_folder 变量为正确的路径。")
        return False
    
    if not os.path.isdir(target_folder):
        print(f"❌ 错误：指定的路径不是文件夹: {target_folder}")
        return False
    
    # 自动获取文件夹中的所有docx文件
    docx_pattern = os.path.join(target_folder, "*.docx")
    file_paths = glob.glob(docx_pattern)
    
    # 过滤掉可能存在的"笔记集合.docx"文件（避免重复处理）
    file_paths = [f for f in file_paths if not os.path.basename(f).startswith("笔记集合")]
    
    if not file_paths:
        print(f"❌ 在文件夹 '{target_folder}' 中未找到任何 .docx 文件")
        return False
    
    # 按文件名排序（确保处理顺序一致）
    file_paths.sort()
    
    print(f"📝 找到 {len(file_paths)} 个docx文件:")
    for i, file_path in enumerate(file_paths, 1):
        print(f"  {i}. {os.path.basename(file_path)}")
    
    # 调用合并函数 (已更新为新函数名)
    success = merge_files_and_overwrite(file_paths)
    
    if success:
        # 输出文件保存在目标文件夹的上级目录
        parent_dir = os.path.dirname(target_folder)
        output_path = os.path.join(parent_dir, "笔记集合.docx")
        print(f"✅ 笔记集合已创建: {output_path}")
    else:
        print("❌ 操作失败或被取消")
    
    return success

def merge_files_and_overwrite(file_paths):
    """
    合并选中的docx文件，如果目标文件已存在则直接覆盖。
    Args:
        file_paths: 选中的文件路径列表
    """
    if not file_paths:
        print("❌ 没有选择文件")
        return False
    
    # 保存到上一层文件夹
    first_file_dir = os.path.dirname(file_paths[0])
    parent_dir = os.path.dirname(first_file_dir)
    output_file = os.path.join(parent_dir, "笔记集合.docx")
    
    print(f"📁 文件所在目录: {first_file_dir}")
    print(f"📍 输出文件路径: {output_file}")
    
    # ⭐ 修改点：不再弹出确认框，如果文件存在，直接提示将要覆盖
    if os.path.exists(output_file):
        print("⚠️ '笔记集合.docx' 已存在，将直接覆盖。")
    
    try:
        # 使用第一个文件作为基础文档
        master_doc = Document(file_paths[0])
        
        # 检查文档中可用的样式
        print("📋 检查可用样式...")
        available_styles = [style.name for style in master_doc.styles]
        print(f"   可用样式: {len(available_styles)} 个")
        
        # 安全地添加集合信息
        first_para = master_doc.paragraphs[0]
        
        # 添加标题
        try:
            title_para = first_para.insert_paragraph_before("📚 笔记集合")
            if 'Title' in available_styles:
                title_para.style = 'Title'
            elif 'Heading 1' in available_styles:
                title_para.style = 'Heading 1'
            else:
                if title_para.runs:
                    title_para.runs[0].font.size = Pt(18)
                    title_para.runs[0].font.bold = True
        except Exception as e:
            print(f"⚠️ 设置标题样式时出错: {str(e)}")
            title_para = first_para.insert_paragraph_before("📚 笔记集合")
        
        # 添加文件信息和创建时间
        current_time = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        file_list_para = first_para.insert_paragraph_before(f"📄 包含文件: {len(file_paths)} 个")
        time_para = first_para.insert_paragraph_before(f"🕒 创建时间: {current_time}")
        first_para.insert_paragraph_before("═" * 60)
        
        # 添加第一个文件的标题
        first_filename = os.path.basename(file_paths[0])
        file_title = first_para.insert_paragraph_before(f"📄 文件来源：{first_filename}")
        
        # 安全地设置文件标题样式
        try:
            if 'Heading 1' in available_styles:
                file_title.style = 'Heading 1'
            else:
                if file_title.runs:
                    file_title.runs[0].font.size = Pt(14)
                    file_title.runs[0].font.bold = True
        except Exception as e:
            print(f"⚠️ 设置文件标题样式时出错: {str(e)}")
        
        # 创建文档合并器
        composer = Composer(master_doc)
        
        # 合并其余选中的文件
        for i, file_path in enumerate(file_paths[1:], 2):
            filename = os.path.basename(file_path)
            print(f"📝 正在合并 ({i}/{len(file_paths)}): {filename}")
            
            try:
                current_doc = Document(file_path)
                current_doc.paragraphs[0].insert_paragraph_before("═" * 60)
                title_para = current_doc.paragraphs[0].insert_paragraph_before(f"📄 文件来源：{filename}")
                
                try:
                    current_available_styles = [style.name for style in current_doc.styles]
                    if 'Heading 1' in current_available_styles:
                        title_para.style = 'Heading 1'
                    else:
                        if title_para.runs:
                            title_para.runs[0].font.size = Pt(14)
                            title_para.runs[0].font.bold = True
                except Exception as e:
                    print(f"⚠️ 设置样式失败: {str(e)}")
                
                if i < len(file_paths):
                    current_doc.add_page_break()
                
                composer.append(current_doc)
                
            except Exception as e:
                print(f"❌ 处理文件 {filename} 时出错: {str(e)}")
                continue
        
        # 保存合并后的文档
        composer.save(output_file)
        print(f"✅ 笔记集合创建成功！")
        print(f"📍 保存位置: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ 合并过程中出现错误: {str(e)}")
        return False

# 运行程序
if __name__ == "__main__":
    print("🚀 启动自动文件夹合并工具...")
    auto_merge_folder_files()